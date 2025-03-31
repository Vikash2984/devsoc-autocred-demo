import os
import asyncio
import pandas as pd
import cloudinary.uploader
from fastapi import FastAPI, File, UploadFile, Form, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from docx import Document
from docx.shared import Pt, RGBColor
from spire.doc import Document as SpireDocument, FileFormat
from dotenv import load_dotenv

load_dotenv()

cloudinary.config(
    cloud_name=os.getenv("CLOUD_NAME"),
    api_key=os.getenv("API_KEY"),
    api_secret=os.getenv("API_SECRET")
)

app = FastAPI()
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])

progress_data = {}

FONT_STYLES = {
    "{name}": Pt(26),
    "{department}": Pt(21),
    "{year}": Pt(21),
    "{event}": Pt(20.5),
    "{date}": Pt(19)
}

def upload_to_cloudinary(file_path, folder):
    result = cloudinary.uploader.upload(file_path, folder=folder, use_filename=True, unique_filename=False, resource_type="raw" if file_path.endswith("_logs.xlsx") else "auto")
    os.remove(file_path)
    return result["secure_url"]

def convert_docx_to_pdf(docx_path, name, email, event, role):
    pdf_path = f"{os.path.dirname(docx_path)}/{name}_{email}_{role}.pdf"
    doc = SpireDocument()
    doc.LoadFromFile(docx_path)
    doc.SaveToFile(pdf_path, FileFormat.PDF)
    os.remove(docx_path)
    return upload_to_cloudinary(pdf_path, f"AutoCred/{event}")

def generate_certificate(template, placeholders, output_folder, event, role):
    os.makedirs(output_folder, exist_ok=True)
    doc = Document(template)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in placeholders.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, "")
                            run = paragraph.add_run(value)
                            run.font.size = FONT_STYLES.get(key, Pt(21.5))
                            run.font.italic = True
                            run.font.color.rgb = RGBColor(171, 124, 52)

    name, email = placeholders['{name}'].strip().replace(" ", "_"), placeholders['{email}'].strip().replace(" ", "_")
    docx_path = f"{output_folder}/{name}_{email}_{role}.docx"
    doc.save(docx_path)

    return convert_docx_to_pdf(docx_path, name, email, event, role)

async def process_bulk_certificates(event_name, event_date, template_path, file_path, pdf_folder, role):
    data = pd.read_excel(file_path, engine="openpyxl").to_dict(orient="records")
    log_file_path = f"{os.path.dirname(file_path)}/{event_name}_logs.xlsx"
    df = pd.DataFrame(columns=["Name", "Email", "Certificate"])
    total_certificates = len(data)

    progress_data[event_name] = {"completed": 0, "total": total_certificates}

    for student in data:
        placeholders = {f'{{{k.lower()}}}': v for k, v in student.items()}
        placeholders.update({'{event}': event_name, '{date}': event_date})
        pdf_url = generate_certificate(template_path, placeholders, pdf_folder, event_name, role)
        df = pd.concat([df, pd.DataFrame([{"Name": student['Name'], "Email": student['Email'], "Certificate": pdf_url}])], ignore_index=True)
        progress_data[event_name]["completed"] += 1  
        await asyncio.sleep(0.1)

    df.to_excel(log_file_path, index=False)
    log_file_url = upload_to_cloudinary(log_file_path, f"AutoCred/{event_name}")
    del progress_data[event_name]

    return log_file_url

@app.get("/progress/{event_name}/total")
async def get_total_certificates(event_name: str):
    return {"total_certificates": progress_data.get(event_name, {}).get("total", "Event not found")}

@app.get("/progress/{event_name}/completed")
async def get_completed_certificates(event_name: str):
    if event_name not in progress_data:
        return JSONResponse({"error": "Event not found"}, status_code=404)

    async def event_stream():
        while event_name in progress_data:
            yield f"data: {progress_data[event_name]['completed']}\n\n"
            await asyncio.sleep(1)

    return StreamingResponse(event_stream(), media_type="text/event-stream")

@app.post("/generate-certificates")
async def generate_certificates(
    background_tasks: BackgroundTasks,
    event_name: str = Form(...), event_date: str = Form(...), template: str = Form(...),
    gen_type: str = Form(...), file: UploadFile | None = File(None), student_name: str = Form(None),
    department: str = Form(None), year: str = Form(None), email: str = Form(None)
):
    script_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = f"{script_dir}/temp{1 if template == 'template1' else 2}.docx"
    role = "participant" if template == "template1" else "organizer"
    pdf_folder = "tmp/Certificates"
    os.makedirs(pdf_folder, exist_ok=True)

    if gen_type == "single":
        placeholders = {'{name}': student_name, '{department}': department, '{year}': year, '{event}': event_name, '{date}': event_date, '{email}': email}
        return JSONResponse({"message": "Certificate generated successfully", "download_url": generate_certificate(template_path, placeholders, pdf_folder, event_name, role)})
    
    elif gen_type == "bulk" and file:
        excel_path = f"{script_dir}/{file.filename}"
        with open(excel_path, "wb") as f:
            f.write(file.file.read())

        log_file_url = await process_bulk_certificates(event_name, event_date, template_path, excel_path, pdf_folder, role)
        return JSONResponse({"message": "Bulk certificate generation completed", "log_file_url": log_file_url})

    return JSONResponse({"error": "Invalid generation type or missing data"}, status_code=400)
